from fastapi import FastAPI, UploadFile, File, Form, HTTPException
from fastapi.responses import StreamingResponse
from fastapi.middleware.cors import CORSMiddleware
import fitz  # PyMuPDF
import docx
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io
import re
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity
import nltk
from nltk.corpus import stopwords
import uvicorn
from google import genai
from google.genai import types
import os
import json
import sqlite3
from datetime import datetime
from dotenv import load_dotenv

# Carregar variáveis de ambiente
load_dotenv()

# Download stopwords
try:
    nltk.data.find('corpora/stopwords')
except LookupError:
    nltk.download('stopwords')

# Configuração do Banco de Dados
DB_PATH = "history.db"

def init_db():
    conn = sqlite3.connect(DB_PATH)
    cursor = conn.cursor()
    cursor.execute('''
        CREATE TABLE IF NOT EXISTS scan_history (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            filename TEXT,
            match_score REAL,
            job_description TEXT,
            resume_text TEXT,
            timestamp DATETIME DEFAULT CURRENT_TIMESTAMP
        )
    ''')
    conn.commit()
    conn.close()

init_db()

# Configuração da IA (Nova Biblioteca Gemini)
GEMINI_API_KEY = os.getenv("GEMINI_API_KEY", "")
client = None
if GEMINI_API_KEY:
    client = genai.Client(api_key=GEMINI_API_KEY)

app = FastAPI(title="DestravaCV API")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

def extract_text_from_pdf(content: bytes) -> str:
    try:
        doc = fitz.open(stream=content, filetype="pdf")
        text = "".join([page.get_text() for page in doc])
        return text
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Erro PDF: {str(e)}")

def extract_text_from_docx(content: bytes) -> str:
    try:
        doc = docx.Document(io.BytesIO(content))
        return "\n".join([para.text for para in doc.paragraphs])
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Erro DOCX: {str(e)}")

def get_keywords(text: str):
    words = set(re.findall(r'\b\w+\b', text.lower()))
    stop_words = set(stopwords.words('portuguese')) | set(stopwords.words('english'))
    return {w for w in words if w not in stop_words and len(w) > 2}

@app.post("/scan")
async def scan_resume(file: UploadFile = File(...), job_description: str = Form(...)):
    try:
        content = await file.read()
        filename = file.filename.lower()
        
        if filename.endswith('.pdf'):
            print(f"Extraindo texto de PDF: {filename}")
            resume_text = extract_text_from_pdf(content)
        elif filename.endswith('.docx'):
            print(f"Extraindo texto de DOCX: {filename}")
            resume_text = extract_text_from_docx(content)
        else:
            print(f"Tentando extrair texto de arquivo desconhecido: {filename}")
            try:
                resume_text = content.decode('utf-8')
            except:
                resume_text = extract_text_from_docx(content)

        print(f"Texto extraído (primeiros 100 caracteres): {resume_text[:100]}")

        if not resume_text.strip():
            print("AVISO: Texto extraído está vazio.")
            resume_text = ""

        # LÓGICA DE MATCH BASEADA EM QUALIFICAÇÕES (IA)
        jd_requirements = []
        if client and job_description.strip():
            try:
                # Extrair apenas os termos técnicos e qualificações essenciais da vaga
                prompt_extract = f"Extraia apenas as Hard Skills (tecnologias, idiomas, certificações) essenciais desta vaga como uma lista simples de termos separados por vírgula. VAGA: {job_description}"
                resp_jd = client.models.generate_content(model='gemini-1.5-flash', contents=prompt_extract)
                jd_requirements = [s.strip().lower() for s in resp_jd.text.split(',')]
            except Exception as ia_err:
                print(f"Erro IA: {ia_err}")
                jd_requirements = list(get_keywords(job_description))
        else:
            jd_requirements = list(get_keywords(job_description))

        # Limpar requisitos vazios ou muito curtos
        jd_requirements = [r for r in jd_requirements if len(r) > 1]
        
        # Verificar quantos requisitos da vaga estão no currículo
        resume_text_lower = resume_text.lower()
        found_requirements = []
        missing_requirements = []

        for req in jd_requirements:
            # Busca simples por palavra-chave (pode ser melhorada com Regex futuramente)
            if req in resume_text_lower:
                found_requirements.append(req)
            else:
                missing_requirements.append(req)

        # Cálculo do Score: Porcentagem de requisitos encontrados
        if not jd_requirements:
            match_score = 0
        else:
            match_score = round((len(found_requirements) / len(jd_requirements)) * 100, 2)

        # Salvar no histórico
        conn = sqlite3.connect(DB_PATH)
        cursor = conn.cursor()
        cursor.execute('''
            INSERT INTO scan_history (filename, match_score, job_description, resume_text)
            VALUES (?, ?, ?, ?)
        ''', (file.filename, match_score, job_description, resume_text))
        conn.commit()
        conn.close()

        return {
            "match_score": match_score,
            "missing_keywords": sorted(missing_requirements)[:20],
            "common_keywords": sorted(found_requirements)[:20],
            "filename": file.filename,
            "resume_text": resume_text
        }
    except Exception as e:
        print(f"Erro no scan: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))

@app.get("/history")
async def get_history():
    conn = sqlite3.connect(DB_PATH)
    conn.row_factory = sqlite3.Row
    cursor = conn.cursor()
    cursor.execute('SELECT * FROM scan_history ORDER BY timestamp DESC LIMIT 10')
    rows = cursor.fetchall()
    history = [dict(row) for row in rows]
    conn.close()
    return history

@app.post("/auto-optimize-resume")
async def auto_optimize_resume(data: dict):
    if not client:
        raise HTTPException(status_code=400, detail="IA não configurada.")

    resume_text = data.get('resume_text', '')
    job_description = data.get('job_description', '')

    prompt = f"""
    Você é um Tech Recruiter Sênior. Reescreva o currículo para máxima aderência à vaga.
    MANTENHA A VERDADE. Foco em RESULTADOS e MÉTRICAS.
    
    RETORNE APENAS JSON VÁLIDO:
    {{
        "name": "nome", "email": "email", "phone": "tel",
        "summary": "resumo",
        "experiences": [{{ "title": "cargo", "company": "empresa", "date": "data", "description": "bullet points" }}],
        "skills": ["skill"]
    }}

    VAGA: {job_description}
    CURRÍCULO: {resume_text}
    """

    try:
        response = client.models.generate_content(
            model='gemini-1.5-flash',
            contents=prompt,
            config=types.GenerateContentConfig(
                response_mime_type='application/json',
            ),
        )
        return json.loads(response.text)
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/generate-cover-letter")
async def generate_cover_letter(data: dict):
    if not client:
        raise HTTPException(status_code=400, detail="IA não configurada.")

    resume_text = data.get('resume_text', '')
    job_description = data.get('job_description', '')

    prompt = f"Escreva uma carta de apresentação persuasiva para esta vaga: {job_description}. Use este currículo como base: {resume_text}. Retorne apenas o texto da carta."

    try:
        response = client.models.generate_content(
            model='gemini-1.5-flash',
            contents=prompt
        )
        return {"cover_letter": response.text.strip()}
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/generate-docx")
async def generate_docx(data: dict):
    try:
        doc = docx.Document()
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(data.get('name', 'Seu Nome').upper())
        run.bold = True
        run.font.size = Pt(18)

        def add_section(title, content):
            p = doc.add_paragraph()
            p.space_before = Pt(15)
            run = p.add_run(title.upper())
            run.bold = True
            run.font.color.rgb = RGBColor(37, 99, 235)
            doc.add_paragraph(content)

        add_section("Resumo Profissional", data.get('summary', ''))
        
        p = doc.add_paragraph()
        p.add_run("EXPERIÊNCIA PROFISSIONAL").bold = True
        for exp in data.get('experiences', []):
            doc.add_paragraph(exp.get('title', '')).bold = True
            doc.add_paragraph(f"{exp.get('company', '')} | {exp.get('date', '')}").italic = True
            doc.add_paragraph(exp.get('description', ''))

        add_section("Habilidades", ", ".join(data.get('skills', [])))

        target = io.BytesIO()
        doc.save(target)
        target.seek(0)
        return StreamingResponse(target, media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

if __name__ == "__main__":
    uvicorn.run(app, host="0.0.0.0", port=8000)
